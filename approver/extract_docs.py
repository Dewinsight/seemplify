"""Extract text from docx and xlsx for rule verification."""
import sys
import os

base = os.path.dirname(os.path.abspath(__file__))

# Docx
try:
    from docx import Document
    doc_path = os.path.join(base, "Sterling_AI_Evaluation_Supplementary_Rules_v2.docx")
    if os.path.exists(doc_path):
        doc = Document(doc_path)
        print("=== DOCX: Sterling_AI_Evaluation_Supplementary_Rules_v2.docx ===\n")
        for para in doc.paragraphs[:80]:  # First 80 paragraphs
            if para.text.strip():
                print(para.text)
        print("\n--- Tables ---")
        for i, table in enumerate(doc.tables[:3]):
            print(f"\nTable {i+1}:")
            for row in table.rows[:15]:
                print(" | ".join(cell.text[:80] for cell in row.cells))
    else:
        print("Docx not found")
except ImportError:
    print("python-docx not installed. Run: pip install python-docx")
except Exception as e:
    print(f"Docx error: {e}")

# Xlsx
try:
    import openpyxl
    xlsx_path = os.path.join(base, "Sterling_AI_Atomic_Rules.xlsx")
    if os.path.exists(xlsx_path):
        wb = openpyxl.load_workbook(xlsx_path, read_only=True, data_only=True)
        print("\n\n=== XLSX: Sterling_AI_Atomic_Rules.xlsx ===\n")
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            print(f"--- Sheet: {sheet_name} ---")
            for row in list(ws.iter_rows(values_only=True))[:100]:
                vals = [str(v)[:60] if v is not None else "" for v in row]
                if any(vals):
                    print(" | ".join(vals))
        wb.close()
    else:
        print("Xlsx not found")
except ImportError:
    print("openpyxl not installed. Run: pip install openpyxl")
except Exception as e:
    print(f"Xlsx error: {e}")
